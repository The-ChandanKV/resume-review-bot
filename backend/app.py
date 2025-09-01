from fastapi import FastAPI, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import pdfplumber
from docx import Document
from io import BytesIO
import google.generativeai as genai
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Gemini API
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
model = genai.GenerativeModel('gemini-pro')

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    messages: List[ChatMessage]

app = FastAPI(
    title="Resume Review Bot API",
    description="API for Resume Review Bot with Gemini AI integration",
    version="1.0.0"
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(file: UploadFile) -> str:
    try:
        with pdfplumber.open(BytesIO(file.file.read())) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

def extract_text_from_docx(file: UploadFile) -> str:
    try:
        doc = Document(BytesIO(file.file.read()))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")

def analyze_resume(text: str) -> Dict[str, Any]:
    # Mock AI analysis - in a real app, this would use an actual AI model
    word_count = len(text.split())
    return {
        "summary": "This is a mock analysis. In a real implementation, this would contain AI-generated feedback.",
        "metrics": {
            "word_count": word_count,
            "readability_score": min(100, max(10, word_count // 50))  # Mock score
        },
        "suggestions": [
            "Add more technical skills",
            "Include quantifiable achievements",
            "Consider adding a professional summary"
        ]
    }

@app.post("/upload_resume")
async def upload_resume(file: UploadFile):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    file_extension = file.filename.split('.')[-1].lower()
    
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file)
    elif file_extension in ['docx', 'doc']:
        text = extract_text_from_docx(file)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX file.")
    
    analysis = analyze_resume(text)
    return {"filename": file.filename, "text": text, "analysis": analysis}

@app.post("/chat")
async def chat(request: ChatRequest):
    try:
        # Convert messages to Gemini format
        chat = model.start_chat(history=[])
        
        # Get the last user message
        last_message = request.messages[-1].content
        
        # Get response from Gemini
        response = chat.send_message(last_message)
        
        return {"response": response.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload_resume")
async def upload_resume(file: UploadFile):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")
    
    try:
        # Extract text based on file type
        file_extension = file.filename.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            with pdfplumber.open(BytesIO(await file.read())) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file_extension in ['docx', 'doc']:
            doc = Document(BytesIO(await file.read()))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        # Get analysis from Gemini
        prompt = f"""
        Analyze this resume and provide feedback in the following format:
        
        **Strengths:**
        - [List key strengths]
        
        **Areas for Improvement:**
        - [List areas for improvement]
        
        **Suggestions:**
        - [List specific suggestions]
        
        Resume:
        {text}
        """
        
        response = model.generate_content(prompt)
        
        return {
            "filename": file.filename,
            "text": text,
            "analysis": response.text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    return {"status": "ok"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
